# -*- coding: utf-8 -*-
from docx import Document
from pathlib import Path
p = Path('盲审专家意见回复_优化版.docx')
doc = Document(p)
print('tables', len(doc.tables), 'paragraphs', len(doc.paragraphs))
for i, para in enumerate(doc.paragraphs[:8], 1):
    if para.text.strip(): print(i, para.text.strip())
for ti, table in enumerate(doc.tables, 1):
    print('TABLE', ti, 'rows', len(table.rows), 'cols', len(table.columns))
    print('HEAD', [c.text.strip().replace('\n',' | ') for c in table.rows[0].cells])
