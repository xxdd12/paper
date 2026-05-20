# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path('盲审专家意见回复_优化版.docx')

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_run_font(run, east='宋体', west='Times New Roman', size=None, bold=None, color=None):
    run.font.name = west
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def add_text_cell(cell, text, size=10.5, bold=False, center=False):
    cell.text = ''
    for idx, para_text in enumerate(text.split('\n')):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        r = p.add_run(para_text)
        set_run_font(r, size=size, bold=bold)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)

def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tcBorders.append(element)
        element.set(qn('w:val'), 'nil')

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True)
    return p

expert1 = [
('1', '论文目录文字采用红色字体，不符合学位论文格式规范要求。', '已按专家意见完成修改。论文已将目录及交叉引用显示颜色统一调整为黑色，保证目录与正文文字风格一致，满足学位论文打印与归档格式要求。\n修改位置：目录页及主文档超链接样式设置。'),
('2', '第一章对立体仓库堆叠货物盘点这一特定场景的痛点描述仍显简略，建议补充更多工业场景下的具体挑战。', '已在第1章“研究背景及研究意义”中补充工业现场约束与任务痛点，包括多层紧密堆叠、相邻箱体遮挡与边界粘连、同类箱体外观相似、固定单视角采集导致三维信息不足、顶层与非顶层结构判别困难，以及现场部署对实时性和轻量化的要求。修改后进一步突出了本文研究的针对性和必要性。\n修改位置：第1章1.1节。'),
('3', '第29页“公式公式（2.15）”中“公式”表述重复，建议删除重复字样，并对全文类似问题进行核查。', '已删除重复表述，并将相关公式引用统一调整为更规范的“式（……）”表述。同时对全文公式引用、图表引用和重复文字进行了检查与修正，以提升论文表达的准确性和规范性。\n修改位置：第2章2.2.3节及全文相关引用表述。'),
('4', '论文中部分页面留白较多，如31页、43页、49页等，建议优化图表位置、段落分页及章节排版。', '已对图表浮动位置、算法环境、段落分页和章节末页排版进行调整，重点优化专家指出的第31页、第43页、第49页附近版式。修改后第31页补充图文衔接内容，第43页补充压缩流程说明和小结内容，第49页由正文自然承接，页面留白明显减少，阅读连续性和版式紧凑性得到改善。\n修改位置：第2章、第3章、第4章相关页面。'),
('5', 'C2SASA、SSConv、C3k2_QDConv等核心改进模块描述较平淡，缺乏与现有同类型轻量化/增强型模块的横向对比讨论。', '已在第2章补充关键模块与同类结构的功能差异对比，从设计动机、结构特点和任务作用等角度说明C2SASA、SSConv和C3k2_QDConv的差异化设计。修改后进一步明确：C2SASA侧重关键区域筛选与注意力增强，SSConv兼顾轻量化下采样与局部结构保持，C3k2_QDConv强化方向性边缘和几何轮廓建模，从而更清晰地体现各模块在分割精度、推理效率和盘点稳定性方面的作用。\n修改位置：第2章2.1.2节，新增关键模块功能差异对比表及相关分析。'),
('6', '实验设计略显单一，与其他实例分割方法的对比较为简单，建议增加系统性对比实验，并可引入公开数据集验证泛化性与鲁棒性。', '已在第4章补充系统性模型对比实验，选取QueryInst、SOLOv2、SparseInst、Mask2Former、YOLOv8s-Seg和YOLO11s-Seg等代表性实例分割方法进行比较，并从box mAP、mask mAP、Params、GFLOPs、DA和CA等指标综合分析。考虑到本文任务依赖up_box/down_box层次标注和layer_count等盘点先验，通用公开数据集难以直接对应本文的库存盘点任务定义，因此本轮修改优先基于真实仓储场景TCSD数据集开展五折交叉验证和多模型对比。关于公开数据集、跨仓库和跨货物类型验证，已在未来工作中作为后续研究方向说明。\n修改位置：第4章4.1节、4.4.2节、4.4.3节及第5章未来工作。'),
('7', '算法部分细节描述不够深入，结构化剪枝与知识蒸馏结合的数学原理和理论支持不足。', '已在第3章补充模型压缩方法的理论分析与公式推导，包括基于BN缩放因子的通道重要性评估、L1稀疏约束训练目标、稀疏项对BN缩放因子的梯度影响、结构化通道剪枝目标、模块结构约束、权重迁移、软目标蒸馏损失与多阶段特征蒸馏损失等内容。同时进一步说明蒸馏损失如何在反向传播过程中对学生模型输出层和中间特征层提供额外约束，从而缓解结构化剪枝后的表达能力下降。\n修改位置：第3章3.2.1节、3.2.3节、3.3.4节。'),
('8', '实验结果部分缺乏详细图表展示不同模型在各项指标下的表现，建议增加性能对比可视化图，并深入分析为什么提出的方法优于对比方法。', '已在第4章完善实验结果展示与讨论，补充不同模型的实例分割性能、盘点准确率、压缩前后性能、推理速度以及不同压缩率设置下的对比表和趋势图。同时增加结果归因分析，指出TCSD场景的主要困难在于相邻箱体边界分离、顶层与非顶层结构判别以及局部遮挡下的轮廓恢复，而YOLO11-StackLite的三个改进模块能够分别对应关键区域筛选、下采样信息保持和方向性几何建模，因此在较低复杂度下取得更稳定的综合性能。\n修改位置：第4章4.4节、4.5节。'),
('9', '近三年的相关文献引用相对偏少，中文文献比例不足，建议补充近几年代表性研究成果，同时适当增加中文文献引用。', '已补充近三年智能仓储、仓储计算机视觉、神经网络压缩和知识蒸馏等方向的代表性研究成果，并增加中文文献引用。新增文献覆盖自动化立体仓库视觉盘点、点云智能盘点、仓储计算机视觉综述、神经网络压缩联合优化和知识蒸馏等内容，使文献综述能够更好反映国内外最新研究进展。\n修改位置：第1章相关研究现状及参考文献列表。'),
('10', '论文部分段落表述重复，影响阅读体验，建议简化语言，避免冗余内容，提高文章流畅性和逻辑性。', '已对全文语言表述进行梳理，重点检查算法改进动机、实验结果分析和章节小结中存在的重复表述，对部分相近内容进行合并、删减和重写。修改后论文结构更加紧凑，章节之间的逻辑衔接更加清晰，文字表达更加简洁规范。\n修改位置：全文，重点为第2章、第3章、第4章及第5章。'),
]

expert2 = [
('1', '第4章消融实验中，C2SASA带来的提升幅度较小，建议加入模块组合的统计显著性检验。', '已在第4章消融实验部分补充五折交叉验证统计结果，并基于相同折次结果增加双侧配对t检验。新增表格报告基线模型、完整模型以及未引入C2SASA的组合模型在box mAP和mask mAP上的均值、标准差和p值。结果表明，完整模型相对于基线模型的提升达到统计显著水平；相对于未引入C2SASA的组合模型，完整模型也保持稳定提升趋势。论文同时说明五折样本量有限，显著性分析主要用于辅助说明模块组合的稳定性，最终结论仍结合复杂度、可视化效果和盘点任务指标综合判断。\n修改位置：第4章4.4.1节，新增消融实验五折统计结果与配对显著性检验表。'),
('2', '第4章模型压缩实验中，蒸馏增强带来的DA和CA提升较小，建议补充多次实验方差以证明方法有效性。', '已在第4章模型压缩实验中补充客观说明。考虑到蒸馏增强相对于剪枝微调模型带来的DA和CA提升幅度较小，论文未将该结果过度解释为强显著性结论，而是将蒸馏定位为压缩后的性能补偿环节。当前论文通过五折平均结果说明蒸馏训练能够在参数量和GFLOPs不变的条件下提供一定性能补偿；同时明确指出，若后续面向更严格统计验证或工程部署，将在相同折次上进行多随机种子重复实验，并报告均值、标准差和显著性检验结果。该处理避免对小幅提升作过强结论，使实验结论更加稳健。\n修改位置：第4章4.5.1节及第5章总结部分。'),
('3', '建议未来在实际边缘计算单元上进一步验证算法性能，并说明当前方法处理不规则堆叠、箱体倾斜等情况的局限性。', '已在第4章和第5章补充边缘部署与方法局限性说明。对于模型压缩部分，论文进一步指出当前实验主要完成参数量、GFLOPs、推理速度和任务指标等代理指标验证，后续将结合Jetson Xavier NX、RK3588等边缘计算平台测试推理时延、资源占用和运行稳定性。对于不规则堆叠、箱体倾斜、严重遮挡或局部缺失等复杂情况，论文已说明当前基于二维掩膜关键点和面积比的估计方法可能产生偏差，后续可结合三维几何信息、多视角图像或视频时序信息进一步提升鲁棒性。\n修改位置：第4章4.5.2节、第5章未来研究展望。'),
]

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.4)
section.right_margin = Cm(2.4)

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
styles['Normal'].font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run('学位论文盲审专家意见回复')
set_run_font(r, size=18, bold=True)

meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = 'Table Grid'
meta.autofit = False
meta_data = [('论文题目', '面向立体仓库堆叠货物盘点的改进实例分割与模型压缩方法研究'), ('学生姓名', '胡文轩'), ('导师姓名', '胡晓敏')]
for i, (k, v) in enumerate(meta_data):
    set_cell_width(meta.rows[i].cells[0], 3.2)
    set_cell_width(meta.rows[i].cells[1], 12.0)
    set_cell_shading(meta.rows[i].cells[0], 'F2F2F2')
    add_text_cell(meta.rows[i].cells[0], k, bold=True)
    add_text_cell(meta.rows[i].cells[1], v)

doc.add_paragraph()
intro = doc.add_paragraph()
r = intro.add_run('根据盲审专家提出的意见和建议，论文已从格式规范、研究背景、算法理论分析、实验验证、结果讨论、参考文献及未来工作等方面进行了认真修改和补充。现将主要修改情况逐条回复如下。')
set_run_font(r, size=10.5)
intro.paragraph_format.first_line_indent = Cm(0.74)
intro.paragraph_format.line_spacing = 1.15
intro.paragraph_format.space_after = Pt(8)

def add_reply_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ['序号', '专家意见', '修改及回复']
    widths = [1.0, 5.7, 8.7]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, 'D9EAF7')
        add_text_cell(cell, h, size=10.5, bold=True, center=True)
    set_repeat_table_header(table.rows[0])
    for no, opinion, reply in rows:
        cells = table.add_row().cells
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)
        add_text_cell(cells[0], no, size=10.5, center=True)
        add_text_cell(cells[1], opinion, size=10.0)
        add_text_cell(cells[2], reply, size=10.0)
    return table

add_section_heading(doc, '一、专家一意见回复')
add_reply_table(doc, expert1)
add_section_heading(doc, '二、专家二意见回复')
add_reply_table(doc, expert2)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(0)
r = p.add_run('说明：以上修改均已体现在当前答辩论文版本中。对涉及后续扩展验证的建议，论文已在结果讨论或未来工作部分作出客观说明，避免对现有实验结论作过度外推。')
set_run_font(r, size=10.5)
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.line_spacing = 1.15

sig = doc.add_table(rows=1, cols=2)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
sig.autofit = False
for i, text in enumerate(['学生签名：', '导师签名：']):
    cell = sig.rows[0].cells[i]
    set_cell_width(cell, 7.5)
    add_text_cell(cell, '\n' + text + '                         年   月   日', size=10.5)
    remove_cell_borders(cell)

doc.save(out)
print(out.resolve())
